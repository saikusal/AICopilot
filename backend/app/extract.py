import io


def extract_resume_text(data: bytes, filename: str, content_type: str | None) -> str:
    """Extract plain text from an uploaded resume (PDF, DOCX, or TXT)."""
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        return _from_pdf(data)
    if name.endswith(".docx") or "word" in ctype or "officedocument" in ctype:
        return _from_docx(data)
    if name.endswith((".txt", ".md")) or ctype.startswith("text/"):
        return data.decode("utf-8", errors="ignore").strip()

    # Unknown type: best-effort decode, otherwise reject.
    text = data.decode("utf-8", errors="ignore").strip()
    if text:
        return text
    raise ValueError("Unsupported file type. Upload a PDF, DOCX, or TXT resume.")


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is not installed") from exc

    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(part.strip() for part in parts if part.strip()).strip()


def _from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc

    document = Document(io.BytesIO(data))
    parts = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
